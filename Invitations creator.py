#!/usr/bin/env python
# coding: utf-8

# In[2]:


import docx, os
os.chdir("/Users/daniellynch/Downloads/")
doc = docx.Document()
doc.add_paragraph('Hello, world!')
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.save('helloworld.docx')


# In[47]:


import docx, os
os.chdir("/Users/daniellynch/Downloads/")
doc = docx.Document()
names = ["Daniel", "Robert", "Tom", "Robocop"]
x = 0
for i in names:
    x+= 1
    para_1 = doc.add_paragraph("It would be the pleasure to have the company of")
    para_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_1.runs[0].underline = True
    para_1.runs[0].bold = True
    para_1.runs[0].font.size = Pt(20)
    
    para_2 = doc.add_paragraph(i)
    para_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_2.runs[0].bold = True
    para_2.runs[0].font.size = Pt(18)
    
    para_3 = doc.add_paragraph("At 74 Brandize House")
    para_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_3.runs[0].italic = True
    para_3.runs[0].font.size = Pt(16)
    
    para_4 = doc.add_paragraph("April the 1st")
    para_4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_4.runs[0].Wingdings = True
    para_4.runs[0].font.size = Pt(14)
    
    para_5 =doc.add_paragraph("At 3pm")
    para_5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_5.runs[0].bold = True
    para_5.runs[0].font.size = Pt(12)
    if x == len(names):
        break
    else:
        doc.paragraphs[-1].runs[-1].add_break(docx.text.run.WD_BREAK.PAGE)

doc.save("Invitations.docx")
print("Invitations Created")


# In[ ]:




