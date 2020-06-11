#!/usr/bin/env python
# coding: utf-8

# In[58]:


import docx, os, openpyxl
os.chdir('/Users/daniellynch/Downloads/')
wb = openpyxl.load_workbook('EmailSheet.xlsx')
ws = wb.active
column_A = []
column_B = []
column_C = []
for cell in ws["A"]:
    column_A.append(cell.value)

for cell in ws["B"]:
    column_B.append(cell.value)

for cell in ws["E"]:
    column_C.append(cell.value)

for a,b,c in zip (column_A[1:5], column_B[1:5], column_C[1:5]):
    doc = docx.Document('Excel information.docx')
    name = (a + "Excel information.docx")
    table = doc.tables[0]
    value_1 = table.cell(0,1).text = (a)
    value_2 = table.cell(1,1).text = (b)
    value_3 = table.cell(2,1).text = (c)

    doc.save(name)
print("Finished transferring")                  
                  
                
            
        
    


# In[37]:


doc = docx.Document('Excel information.docx')
print(doc.tables)
table = doc.tables[0]
table.cell(0,1).text = ("Email: " + "This is a test")
doc.save("Test Run.docx")


# In[ ]:




